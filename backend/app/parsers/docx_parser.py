"""DOCX parser — extracts the first table found in the document."""

from __future__ import annotations

from pathlib import Path
from typing import Any


def parse_docx(path: Path, *, max_preview_rows: int = 25) -> tuple[list[str], list[dict[str, Any]], int]:
    from docx import Document

    doc = Document(str(path))
    if not doc.tables:
        return [], [], 0

    table = doc.tables[0]
    raw_rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    if not raw_rows:
        return [], [], 0

    headers = [c or f"col_{i}" for i, c in enumerate(raw_rows[0])]
    records: list[dict[str, Any]] = []
    for row in raw_rows[1:]:
        rec = {h: row[i] if i < len(row) else "" for i, h in enumerate(headers)}
        records.append(rec)
    return headers, records[:max_preview_rows], len(records)


def iter_docx_rows(path: Path) -> tuple[list[str], list[dict[str, Any]]]:
    headers, _, _ = parse_docx(path, max_preview_rows=10**9)
    h2, rows, _ = parse_docx(path, max_preview_rows=10**9)
    return headers or h2, rows
